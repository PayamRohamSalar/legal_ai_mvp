# src/data_processing/document_processor.py
"""
Document processor module for extracting and processing legal documents.
Handles Word documents and prepares them for vectorization.
"""

import re
from typing import List, Dict, Optional, Any
from pathlib import Path
import logging
from dataclasses import dataclass
from datetime import datetime

import docx
from docx import Document as DocxDocument

# Configure logging
logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Container for processed document data."""
    filename: str
    content: str
    raw_text: str
    metadata: Dict[str, Any]
    processed_at: datetime
    document_type: str = "legal"
    encoding: str = "utf-8"


class DocumentProcessor:
    """
    Main document processor class for handling legal documents.
    Supports extraction from various formats with focus on Word documents.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = ['.docx', '.doc', '.txt']
        logger.info("DocumentProcessor initialized")
    
    def process_file(self, file_path: Path) -> ProcessedDocument:
        """
        Process a single document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument object containing extracted content and metadata
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        logger.info(f"Processing file: {file_path}")
        
        # Extract content based on file type
        if file_path.suffix.lower() == '.docx':
            content, raw_text = self._extract_from_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            content, raw_text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Handler not implemented for: {file_path.suffix}")
        
        # Extract basic metadata
        metadata = self._extract_file_metadata(file_path)
        
        # Create processed document
        processed_doc = ProcessedDocument(
            filename=file_path.name,
            content=content,
            raw_text=raw_text,
            metadata=metadata,
            processed_at=datetime.now()
        )
        
        logger.info(f"Successfully processed {file_path.name}")
        return processed_doc
    
    def _extract_from_docx(self, file_path: Path) -> tuple[str, str]:
        """
        Extract text content from a Word document.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Tuple of (structured_content, raw_text)
        """
        try:
            doc = docx.Document(str(file_path))
            
            # Extract all paragraphs
            paragraphs = []
            raw_paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Keep paragraph structure
                    paragraphs.append(text)
                    raw_paragraphs.append(text)
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_line = " | ".join(row_text)
                        paragraphs.append(table_line)
                        raw_paragraphs.append(table_line)
            
            # Join with double newlines to preserve structure
            structured_content = "\n\n".join(paragraphs)
            raw_text = "\n".join(raw_paragraphs)
            
            return structured_content, raw_text
            
        except Exception as e:
            logger.error(f"Error extracting from docx: {e}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> tuple[str, str]:
        """
        Extract text content from a plain text file.
        
        Args:
            file_path: Path to the .txt file
            
        Returns:
            Tuple of (structured_content, raw_text)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # For text files, structured and raw are the same
            return content, content
            
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='windows-1256') as f:
                content = f.read()
            return content, content
    
    def _extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file metadata
        """
        stat = file_path.stat()
        
        metadata = {
            'file_name': file_path.name,
            'file_size': stat.st_size,
            'file_extension': file_path.suffix,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'file_path': str(file_path.absolute())
        }
        
        return metadata
    
    def process_directory(self, directory_path: Path) -> List[ProcessedDocument]:
        """
        Process all documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of ProcessedDocument objects
        """
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory path: {directory_path}")
        
        processed_docs = []
        
        # Find all supported files
        for ext in self.supported_formats:
            for file_path in directory_path.glob(f"*{ext}"):
                try:
                    doc = self.process_file(file_path)
                    processed_docs.append(doc)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        logger.info(f"Processed {len(processed_docs)} documents from {directory_path}")
        return processed_docs
    
    def validate_document(self, doc: ProcessedDocument) -> bool:
        """
        Validate that a document has been properly processed.
        
        Args:
            doc: ProcessedDocument to validate
            
        Returns:
            True if valid, False otherwise
        """
        if not doc.content or len(doc.content.strip()) == 0:
            logger.warning(f"Document {doc.filename} has no content")
            return False
        
        if not doc.metadata:
            logger.warning(f"Document {doc.filename} has no metadata")
            return False
        
        return True